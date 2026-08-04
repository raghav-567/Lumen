"""DOCX document parser using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.table_linearizer import ParsedTable

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> tuple[str, int, list[ParsedTable]]:
    """Parse a DOCX file and return (text, page_count_estimate, tables).

    ``doc.paragraphs`` excludes table cell text, so tables are extracted
    separately from ``doc.tables`` — otherwise every table in a Word document
    (a common home for policy targets/limits) would be silently dropped.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    tables: list[ParsedTable] = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows and any(any(c.strip() for c in r) for r in rows):
            tables.append(ParsedTable(rows=rows))

    # Estimate page count (rough: ~3000 chars per page)
    page_count = max(1, len(full_text) // 3000)

    logger.info(
        f"Parsed DOCX {path.name}: ~{page_count} pages, {len(full_text)} chars, "
        f"{len(tables)} tables"
    )
    return full_text, page_count, tables
